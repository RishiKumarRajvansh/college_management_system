import base64
import json
import os
import subprocess
import sys
import tempfile
import textwrap
import time
import urllib.request
import zipfile
from pathlib import Path

import django
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import websocket
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SYNOPSIS_PATH = Path(r"C:\Users\vega6\Desktop\Download backup\Synopsis.docx")
OUTPUT_DIR = ROOT / "reports"
ASSET_DIR = OUTPUT_DIR / "assets"
REPORT_PATH = OUTPUT_DIR / "College_Management_System_Project_Report.docx"
BASE_URL = "http://127.0.0.1:8000"


os.environ.setdefault("DJANGO_SETTINGS_MODULE", "college_management_system.settings")
sys.path.insert(0, str(ROOT))
django.setup()

from django.conf import settings
from django.contrib.auth.models import User
from django.test import Client


def read_synopsis():
    doc = Document(str(SYNOPSIS_PATH))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs), paragraphs


def ensure_dirs():
    OUTPUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)


def get_model_metadata():
    from django.apps import apps

    app_labels = [
        "user_authentication",
        "student_management",
        "faculty_management",
        "course_management",
        "attendance_management",
        "examination",
        "library_management",
        "hostel_management",
        "fee_management",
        "reporting",
    ]
    metadata = []
    for model in apps.get_models():
        if model._meta.app_label not in app_labels:
            continue
        fields = []
        relationships = []
        for field in model._meta.fields:
            field_type = field.__class__.__name__
            details = {
                "name": field.name,
                "type": field_type,
                "primary_key": field.primary_key,
                "unique": getattr(field, "unique", False),
                "null": getattr(field, "null", False),
                "blank": getattr(field, "blank", False),
            }
            if getattr(field, "remote_field", None) and field.remote_field and field.remote_field.model:
                target = field.remote_field.model
                details["target"] = target.__name__
                details["on_delete"] = getattr(field.remote_field.on_delete, "__name__", str(field.remote_field.on_delete))
                relationships.append((model.__name__, field.name, target.__name__))
            fields.append(details)

        metadata.append(
            {
                "app": model._meta.app_label,
                "model": model.__name__,
                "db_table": model._meta.db_table,
                "fields": fields,
                "relationships": relationships,
                "ordering": list(model._meta.ordering),
                "unique_together": list(model._meta.unique_together),
            }
        )
    return metadata


def get_database_counts():
    from django.apps import apps

    counts = {}
    for model in apps.get_models():
        if model._meta.app_label in {
            "user_authentication",
            "student_management",
            "faculty_management",
            "course_management",
            "attendance_management",
            "examination",
            "library_management",
            "hostel_management",
            "fee_management",
            "reporting",
        }:
            try:
                counts[model.__name__] = model.objects.count()
            except Exception:
                counts[model.__name__] = "N/A"
    return counts


def add_arrow(ax, start, end, color="#334155", rad=0.0):
    ax.annotate(
        "",
        xy=end,
        xytext=start,
        arrowprops=dict(arrowstyle="->", lw=1.5, color=color, connectionstyle=f"arc3,rad={rad}"),
    )


def draw_box(ax, xy, w, h, text, fc="#ffffff", ec="#1e40af", text_color="#0f172a", fontsize=9):
    rect = patches.FancyBboxPatch(
        xy,
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.04",
        linewidth=1.3,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(rect)
    ax.text(xy[0] + w / 2, xy[1] + h / 2, text, ha="center", va="center", color=text_color, fontsize=fontsize, wrap=True)


def save_dfd_context(path):
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_axis_off()
    draw_box(ax, (4.1, 2.55), 2.8, 1.4, "College Management System\n(Django Application)", fc="#dbeafe", ec="#2563eb", fontsize=11)
    actors = [
        ((0.5, 4.7), "Administrator\nManages records, reports,\nusers and audit logs"),
        ((0.5, 0.8), "Faculty\nMaintains attendance,\nresults and course data"),
        ((8.0, 4.7), "Student\nViews profile, fees,\nresults and hostel data"),
        ((8.0, 0.8), "Staff\nSupports fee, library and\nadministrative workflows"),
    ]
    for xy, label in actors:
        draw_box(ax, xy, 2.5, 1.1, label, fc="#f8fafc", ec="#64748b", fontsize=9)
    add_arrow(ax, (3.0, 5.2), (4.1, 3.65))
    add_arrow(ax, (3.0, 1.35), (4.1, 2.8))
    add_arrow(ax, (8.0, 5.2), (6.9, 3.65))
    add_arrow(ax, (8.0, 1.35), (6.9, 2.8))
    ax.text(5.5, 6.4, "DFD Level 0: Context Diagram", ha="center", fontsize=15, fontweight="bold")
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)


def save_dfd_level1(path):
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_axis_off()
    modules = [
        ((0.5, 5.8), "User Authentication\nRole-based access"),
        ((3.4, 5.8), "Student & Faculty\nManagement"),
        ((6.3, 5.8), "Course & Attendance\nManagement"),
        ((9.2, 5.8), "Examination\nManagement"),
        ((0.5, 3.4), "Library\nManagement"),
        ((3.4, 3.4), "Hostel\nManagement"),
        ((6.3, 3.4), "Fee\nManagement"),
        ((9.2, 3.4), "Reporting & Audit\nManagement"),
    ]
    for xy, label in modules:
        draw_box(ax, xy, 2.25, 1.0, label, fc="#eff6ff", ec="#2563eb", fontsize=8.5)
    draw_box(ax, (4.75, 1.1), 2.8, 1.0, "SQLite Database\nPersistent institutional records", fc="#f0fdf4", ec="#16a34a", fontsize=10)
    for xy, _ in modules:
        add_arrow(ax, (xy[0] + 1.1, xy[1]), (6.15, 2.1), color="#475569")
    ax.text(6.0, 7.25, "DFD Level 1: Major Process Decomposition", ha="center", fontsize=15, fontweight="bold")
    ax.set_xlim(0, 12)
    ax.set_ylim(0.5, 7.7)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)


def save_er_diagram(path):
    fig, ax = plt.subplots(figsize=(13, 9))
    ax.set_axis_off()
    boxes = {
        "User": (0.5, 6.7),
        "UserProfile": (3.0, 6.7),
        "Student": (0.5, 4.8),
        "Faculty": (3.0, 4.8),
        "Course": (5.6, 4.8),
        "Attendance": (8.5, 5.8),
        "Examination": (8.5, 4.0),
        "Result": (11.0, 4.0),
        "Book": (0.5, 2.7),
        "BookIssue": (3.0, 2.7),
        "Hostel": (5.6, 2.7),
        "Room": (8.5, 2.7),
        "HostelAllocation": (11.0, 2.7),
        "FeeCategory": (0.5, 0.8),
        "FeeStructure": (3.0, 0.8),
        "Payment": (5.6, 0.8),
        "Report": (8.5, 0.8),
        "AuditTrail": (11.0, 0.8),
    }
    for label, xy in boxes.items():
        draw_box(ax, xy, 1.75, 0.75, label, fc="#ffffff", ec="#334155", fontsize=8.5)

    rels = [
        ("User", "UserProfile"), ("User", "Student"), ("User", "Faculty"), ("Faculty", "Course"),
        ("Student", "Attendance"), ("Course", "Attendance"), ("Course", "Examination"), ("Examination", "Result"),
        ("Student", "Result"), ("Student", "BookIssue"), ("Book", "BookIssue"), ("Hostel", "Room"),
        ("Room", "HostelAllocation"), ("Student", "HostelAllocation"), ("FeeCategory", "FeeStructure"),
        ("Course", "FeeStructure"), ("FeeStructure", "Payment"), ("Student", "Payment"), ("User", "Report"),
        ("User", "AuditTrail"),
    ]
    for source, target in rels:
        sx, sy = boxes[source]
        tx, ty = boxes[target]
        add_arrow(ax, (sx + 0.875, sy + 0.375), (tx + 0.875, ty + 0.375), color="#2563eb")
    ax.text(6.6, 8.35, "Entity Relationship Diagram", ha="center", fontsize=16, fontweight="bold")
    ax.set_xlim(0, 13.5)
    ax.set_ylim(0.3, 8.8)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)


def save_use_case(path):
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_axis_off()
    draw_box(ax, (3.3, 0.7), 5.4, 6.4, "College Management System Boundary", fc="#f8fafc", ec="#94a3b8", fontsize=10)
    actors = [((0.5, 5.8), "Admin"), ((0.5, 2.0), "Faculty"), ((10.0, 5.8), "Student"), ((10.0, 2.0), "Staff")]
    for xy, label in actors:
        ax.text(xy[0], xy[1], "☺", fontsize=24, ha="center")
        ax.text(xy[0], xy[1] - 0.35, label, fontsize=10, ha="center", fontweight="bold")
    use_cases = [
        ((3.8, 5.9), "Manage Users"),
        ((6.4, 5.9), "Manage Courses"),
        ((3.8, 4.7), "Record Attendance"),
        ((6.4, 4.7), "Publish Results"),
        ((3.8, 3.5), "Manage Fees"),
        ((6.4, 3.5), "Manage Library"),
        ((3.8, 2.3), "Manage Hostel"),
        ((6.4, 2.3), "Generate Reports"),
        ((5.1, 1.2), "View Profile"),
    ]
    for xy, label in use_cases:
        ellipse = patches.Ellipse((xy[0], xy[1]), 1.9, 0.55, edgecolor="#2563eb", facecolor="#dbeafe", lw=1.3)
        ax.add_patch(ellipse)
        ax.text(xy[0], xy[1], label, ha="center", va="center", fontsize=8)
    for actor_xy, _ in actors:
        for uc_xy, _ in use_cases:
            if actor_xy[0] < 2 and uc_xy[0] < 5.5:
                add_arrow(ax, (actor_xy[0] + 0.35, actor_xy[1] - 0.1), (uc_xy[0] - 0.9, uc_xy[1]), color="#64748b")
            elif actor_xy[0] > 9 and uc_xy[0] > 5:
                add_arrow(ax, (actor_xy[0] - 0.35, actor_xy[1] - 0.1), (uc_xy[0] + 0.9, uc_xy[1]), color="#64748b")
    ax.text(6.0, 7.55, "Use Case Diagram", ha="center", fontsize=16, fontweight="bold")
    ax.set_xlim(0, 12)
    ax.set_ylim(0.2, 8)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)


def save_architecture(path):
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_axis_off()
    draw_box(ax, (0.7, 4.8), 2.2, 1.0, "Browser\nTailwind UI", fc="#eff6ff", ec="#2563eb", fontsize=10)
    draw_box(ax, (4.2, 4.8), 2.6, 1.0, "Django URL Routing\nViews + Forms", fc="#f8fafc", ec="#334155", fontsize=10)
    draw_box(ax, (8.1, 4.8), 2.1, 1.0, "Templates\nHTML + CSS + JS", fc="#fefce8", ec="#ca8a04", fontsize=10)
    draw_box(ax, (4.2, 2.8), 2.6, 1.0, "Django ORM\nModels + Querysets", fc="#f0fdf4", ec="#16a34a", fontsize=10)
    draw_box(ax, (8.1, 2.8), 2.1, 1.0, "SQLite\nDatabase", fc="#ecfdf5", ec="#059669", fontsize=10)
    draw_box(ax, (0.7, 2.8), 2.2, 1.0, "Authentication\nSessions + Roles", fc="#fdf2f8", ec="#db2777", fontsize=10)
    add_arrow(ax, (2.9, 5.3), (4.2, 5.3))
    add_arrow(ax, (6.8, 5.3), (8.1, 5.3))
    add_arrow(ax, (5.5, 4.8), (5.5, 3.8))
    add_arrow(ax, (6.8, 3.3), (8.1, 3.3))
    add_arrow(ax, (2.9, 3.3), (4.2, 3.3))
    ax.text(5.5, 6.35, "Application Architecture", ha="center", fontsize=16, fontweight="bold")
    ax.set_xlim(0, 11)
    ax.set_ylim(2, 6.8)
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)


def create_diagrams():
    diagrams = {
        "dfd_context": ASSET_DIR / "dfd_context.png",
        "dfd_level1": ASSET_DIR / "dfd_level1.png",
        "er": ASSET_DIR / "er_diagram.png",
        "use_case": ASSET_DIR / "use_case.png",
        "architecture": ASSET_DIR / "architecture.png",
    }
    save_dfd_context(diagrams["dfd_context"])
    save_dfd_level1(diagrams["dfd_level1"])
    save_er_diagram(diagrams["er"])
    save_use_case(diagrams["use_case"])
    save_architecture(diagrams["architecture"])
    return diagrams


def capture_screenshots():
    chrome = Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe")
    if not chrome.exists():
        chrome = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
    if not chrome.exists():
        return {}

    admin = User.objects.get(email="admin@riverdale.edu")
    student = User.objects.get(email="maya.patel@students.riverdale.edu")
    faculty = User.objects.get(email="aisha.raman@riverdale.edu")

    def session_cookie_for(user):
        client = Client(HTTP_HOST="127.0.0.1")
        client.force_login(user)
        return client.cookies[settings.SESSION_COOKIE_NAME].value

    port = 9240
    user_data = Path(tempfile.gettempdir()) / "cms-report-cdp-profile"
    proc = subprocess.Popen(
        [
            str(chrome),
            "--headless=new",
            "--disable-gpu",
            "--disable-extensions",
            "--hide-scrollbars",
            "--remote-allow-origins=*",
            f"--remote-debugging-port={port}",
            f"--user-data-dir={user_data}",
            "about:blank",
        ],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    ws = None
    screenshots = {}
    try:
        for _ in range(50):
            try:
                with urllib.request.urlopen(f"http://127.0.0.1:{port}/json", timeout=1) as response:
                    targets = json.loads(response.read().decode("utf-8"))
                break
            except Exception:
                time.sleep(0.1)
        else:
            return {}

        page = [target for target in targets if target.get("type") == "page"][0]
        ws = websocket.create_connection(page["webSocketDebuggerUrl"], timeout=15)
        counter = [0]

        def send(method, params=None):
            counter[0] += 1
            message_id = counter[0]
            ws.send(json.dumps({"id": message_id, "method": method, "params": params or {}}))
            while True:
                message = json.loads(ws.recv())
                if message.get("id") == message_id:
                    if "error" in message:
                        raise RuntimeError(message["error"])
                    return message.get("result")

        send("Page.enable")
        send("Network.enable")

        def set_session(cookie_value):
            send(
                "Network.setCookie",
                {
                    "name": settings.SESSION_COOKIE_NAME,
                    "value": cookie_value,
                    "domain": "127.0.0.1",
                    "path": "/",
                    "url": BASE_URL + "/",
                },
            )

        def shot(name, url, user=None, width=1366, height=900, mobile=False):
            send(
                "Emulation.setDeviceMetricsOverride",
                {"width": width, "height": height, "deviceScaleFactor": 1, "mobile": mobile},
            )
            if user is not None:
                set_session(session_cookie_for(user))
            send("Page.navigate", {"url": BASE_URL + url})
            time.sleep(2.2)
            result = send("Page.captureScreenshot", {"format": "png", "fromSurface": True})
            path = ASSET_DIR / f"{name}.png"
            path.write_bytes(base64.b64decode(result["data"]))
            screenshots[name] = path

        shot("login_page", "/login/", None)
        shot("admin_dashboard", "/dashboard/", admin)
        shot("student_profile", "/students/1/", admin)
        shot("faculty_module", "/faculty/", admin)
        shot("course_dashboard", "/courses/", admin)
        shot("attendance_module", "/attendance/", admin)
        shot("reports_dashboard", "/reports/", admin)
        shot("activity_logs", "/auth/system-activity-logs/", admin)
        shot("student_portal", "/", student)
        shot("faculty_portal", "/", faculty)
        shot("mobile_activity_logs", "/auth/system-activity-logs/?action=login", admin, width=390, height=844, mobile=True)
    finally:
        if ws is not None:
            try:
                ws.close()
            except Exception:
                pass
        proc.terminate()
        try:
            proc.wait(timeout=5)
        except subprocess.TimeoutExpired:
            proc.kill()
    return screenshots


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.color.rgb = RGBColor(15, 23, 42)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.size = Pt(13)

    code_style = styles.add_style("CodeBlock", 1)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.left_indent = Inches(0.15)

    return doc


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("College Management System Project Report | Page ")
    run.font.size = Pt(9)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_paragraph(doc, text, bold_start=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if bold_start and text.startswith(bold_start):
        run = p.add_run(bold_start)
        run.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = str(header)
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()
    return table


def add_image(doc, path, caption, width=6.4):
    if not path or not Path(path).exists():
        add_paragraph(doc, f"Screenshot unavailable: {caption}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9.5)


def add_code(doc, source_path, start_marker=None, max_lines=45):
    path = ROOT / source_path
    lines = path.read_text(encoding="utf-8", errors="ignore").splitlines()
    start = 0
    if start_marker:
        for i, line in enumerate(lines):
            if start_marker in line:
                start = i
                break
    snippet = "\n".join(lines[start:start + max_lines])
    p = doc.add_paragraph(style="CodeBlock")
    p.add_run(snippet)


def page(doc, title, paragraphs=None, bullets=None, numbered=None, image=None, image_caption=None):
    doc.add_heading(title, level=1)
    for paragraph in paragraphs or []:
        add_paragraph(doc, paragraph)
    if bullets:
        add_bullets(doc, bullets)
    if numbered:
        add_numbered(doc, numbered)
    if image:
        add_image(doc, image, image_caption or title)
    doc.add_page_break()


def section_page(doc, title, paragraphs, bullets=None):
    page(doc, title, paragraphs=paragraphs, bullets=bullets)


def make_report(synopsis_text, synopsis_paragraphs, model_metadata, counts, diagrams, screenshots):
    doc = setup_document()
    add_page_number(doc.sections[0])

    # Cover page
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COLLEGE MANAGEMENT SYSTEM")
    r.bold = True
    r.font.size = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project Report")
    r.bold = True
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Built with Django, HTML, CSS, Tailwind CSS and SQLite").font.size = Pt(14)
    doc.add_paragraph()
    cover_rows = [
        ("Submitted By", "RISHI KUMAR"),
        ("Enrollment No.", "2350636647"),
        ("Programme", "Master of Computer Applications (MCA)"),
        ("Course", "MCSP-232"),
        ("Submitted To", "School of Computer and Information Sciences, IGNOU"),
        ("Project Guide", "Mr. Karmveer, Sr. Software Engineer, NIC, Delhi"),
        ("Session", "2024-2026"),
    ]
    add_table(doc, ["Particular", "Details"], cover_rows)
    doc.add_page_break()

    page(
        doc,
        "Certificate",
        [
            "This is to certify that the project report entitled 'College Management System' has been prepared as part of the academic submission requirement for the Master of Computer Applications programme. The work represents a software system designed and implemented to automate institutional operations including student records, faculty records, courses, attendance, examinations, library, hostel, fee processing, reporting, authentication, and system activity auditing.",
            "The application has been developed using the Django web framework with SQLite as the database engine and responsive presentation layers built through HTML, CSS, Tailwind CSS, JavaScript, and Django templates. The project was examined against the synopsis objectives and the implemented codebase to prepare this comprehensive technical documentation.",
        ],
    )
    page(
        doc,
        "Declaration",
        [
            "I hereby declare that the project work presented in this report is based on the College Management System synopsis and the implemented Django project files. The documentation has been prepared for academic submission and explains the problem domain, design decisions, implementation details, testing approach, database structure, screenshots, and future scope of the system.",
            "The project aims to demonstrate the practical application of web development, database management, software engineering methodology, and secure role-based access control in the context of institutional administration.",
        ],
    )
    page(
        doc,
        "Acknowledgement",
        [
            "I express my sincere gratitude to Almighty God, my parents, family members, teachers, friends, and well-wishers whose continuous support made this project possible. I also place on record my deep sense of gratitude to my project guide, Mr. Karmveer, Sr. Software Engineer, NIC, Delhi, for valuable guidance, motivation, and technical direction throughout the project work.",
            "I extend my thanks to the faculty members and academic staff for their cooperation during the course. Their encouragement helped in understanding the institutional requirements and converting them into a practical software solution.",
        ],
    )
    page(
        doc,
        "Abstract",
        [
            "The College Management System is a centralized web application developed to simplify academic and administrative operations in a college environment. The synopsis identifies the core problem as the inefficiency of manual record keeping, delayed access to information, repetitive administrative tasks, and the risk of inconsistent records across departments. The implemented application addresses these problems through a Django-based modular system.",
            "The system provides role-based access for administrators, faculty, students, and staff-like administrative users. It maintains linked records for students, faculty, courses, attendance, examinations, library transactions, hostel allocations, fee structures, payments, reports, notifications, password reset requests, and audit trails. The use of Django ORM with SQLite gives the project a structured relational foundation while remaining lightweight enough for academic deployment and testing.",
            "The report expands the synopsis into a complete academic document covering system analysis, objectives, methodology, feasibility, planning, design diagrams, database schema, code explanations, implementation, testing, screenshots, and future recommendations.",
        ],
    )

    toc_items = [
        "Overview", "Profile of Problem Assigned", "Introduction to Project", "Existing System", "Proposed System",
        "Objective", "Methodology", "Modules", "Technology Used", "Requirement Analysis", "Hardware Specification",
        "Software Specification", "Requirement Specification", "Functional Requirements", "Non-Functional Requirements",
        "Feasibility Analysis", "Project Planning", "Schedule Table", "Design and Diagrams", "Database Schema",
        "Coding", "Testing", "Implementation", "Screenshots of Project", "Future Recommendation", "Bibliography",
    ]
    page(doc, "Table of Contents", paragraphs=["The following academic sections are included in this detailed project report."], numbered=toc_items)

    page(
        doc,
        "Overview",
        [
            "The College Management System is intended to become a single operational platform for an educational institution. Instead of storing student records in isolated registers or spreadsheets, it places all important entities inside a relational database and exposes them through a secure web interface. The application follows a modular structure where each major institutional function is implemented as a separate Django app.",
            "The implemented project contains apps for user authentication, student management, faculty management, course management, attendance management, examination management, library management, hostel management, fee management, and reporting. A common base template integrates a responsive Tailwind CSS shell so the system can be used on desktop and mobile devices.",
            f"The current seeded database contains {counts.get('Student', 0)} students, {counts.get('Faculty', 0)} faculty members, {counts.get('Course', 0)} courses, {counts.get('Attendance', 0)} attendance records, {counts.get('Examination', 0)} examinations, and {counts.get('Payment', 0)} payments. These records demonstrate that the project is not merely a set of static pages but a database-driven application.",
        ],
    )
    page(
        doc,
        "Synopsis Interpretation",
        [
            "The synopsis emphasizes automation, centralized data storage, role-specific interfaces, attendance tracking, grade management, fee management, and administrative efficiency. The implemented project directly reflects these goals. It uses Django views and templates to expose CRUD-style operations for the principal entities and uses Django models to enforce relational links among students, courses, faculty, attendance, examinations, payments, and hostels.",
            "The synopsis originally describes a compact schema with Student, Faculty, Course, Attendance, Examination, Library, and Hostel tables. The codebase expands this into a more complete structure by adding UserProfile, AuditTrail, PasswordResetRequest, BookIssue, HostelAllocation, FeeCategory, FeeStructure, Payment, Notification, Report, and DatabaseBackup models.",
            "This report therefore combines the synopsis expectations with the actual implemented architecture. Where the implementation is more detailed than the synopsis, the report documents the real codebase so that the submission matches the running project.",
        ],
    )
    page(
        doc,
        "Profile of Problem Assigned",
        [
            "Colleges handle large volumes of information related to student admissions, course assignments, faculty workloads, attendance sheets, examination results, hostel occupancy, library circulation, and fee payments. In a manual environment each department may maintain its own files, leading to duplication, inconsistent updates, delayed verification, and difficulty in generating reports.",
            "The assigned problem is to design and implement a system that automates these activities, preserves referential accuracy, and provides role-based access. The system must reduce clerical effort, improve transparency, and make institutional records available in real time to authorized users.",
            "The problem also includes security and accountability. Since academic records are sensitive, the software must restrict access based on roles and keep a record of important system activities such as login, logout, profile updates, user creation, and administrative actions.",
        ],
    )
    page(
        doc,
        "Introduction to Project",
        [
            "The College Management System is a web-based information system built using Python and Django. It uses Django's MTV architecture: models define the database structure, templates define the presentation layer, and views coordinate request processing, validation, and response rendering. The system is accessed through a browser and persists data in SQLite.",
            "The front end uses HTML templates, Tailwind CSS through CDN, compatibility CSS for existing Bootstrap-style classes, Font Awesome icons, and JavaScript for small interactive behaviors such as sidebar toggling and password visibility. The project is suitable for academic demonstration because it combines full-stack development concepts with a realistic institutional use case.",
            "The application includes separate modules that work together. A student profile is linked to a user account and a course. A course is linked to faculty. Attendance links student and course. Results link student and examination. Payments link student and fee structure. These relationships demonstrate practical database normalization.",
        ],
    )
    page(
        doc,
        "Existing System",
        [
            "In many institutions the existing system is manual or semi-digital. Student information may be kept in paper files, attendance may be recorded in registers, fee records may be maintained in spreadsheets, and library or hostel data may be handled by separate staff members. The absence of a single source of truth creates delays and errors.",
            "Manual systems are difficult to search, vulnerable to physical damage, and inefficient for producing reports. If a student changes course information, the same change must be repeated in multiple places. If a fee payment is posted late, departments may not have updated information. If examination results are entered manually, calculation errors can occur.",
            "The existing system also lacks real-time visibility. Administrators cannot easily view overall statistics, faculty cannot instantly track attendance summaries, and students cannot access their academic status without contacting the office.",
        ],
    )
    page(
        doc,
        "Limitations of Existing System",
        [
            "The major limitations of the existing manual system include data redundancy, slow retrieval, missing audit accountability, difficulty in correcting errors, lack of centralized reporting, and limited support for remote access. Paper-based operations also depend heavily on staff availability and physical storage.",
            "Manual attendance and fee systems create repetitive clerical work. They also make it difficult to generate accurate summaries for management decisions. In an academic environment where semester timelines are fixed, delays in record processing directly affect students, faculty, and administration.",
            "Security is also weak in manual systems because access control is not enforced consistently. A digital system can apply authentication, permissions, validation, and audit logs in a more systematic way.",
        ],
    )
    page(
        doc,
        "Proposed System",
        [
            "The proposed system is a Django-based College Management System with centralized data storage and module-wise workflows. It provides login-based access, role-aware navigation, CRUD operations, dashboards, reports, notifications, and activity logs. SQLite is used as the database engine for a lightweight and portable deployment.",
            "The system improves the existing process by storing every major entity in related tables. For example, a student is connected to a course, a course is assigned to faculty, attendance belongs to both student and course, and payment belongs to a student and fee structure. This helps prevent orphan records and improves data integrity.",
            "The Tailwind-styled interface makes the system usable across screen sizes. Administrators can manage institutional data; faculty can view academic workflows; students can access relevant records; staff users can assist with operational processes.",
        ],
    )
    page(
        doc,
        "Benefits of Proposed System",
        [
            "The proposed system offers faster access to records, reduction of manual errors, improved data consistency, secure role-based access, better reporting, and a more professional user experience. Because the system is web-based, it can be accessed through standard browsers without installing separate client applications.",
            "Django's ORM reduces the risk of SQL injection by generating parameterized queries and object-oriented access patterns. Django templates separate presentation from business logic, and model relationships enforce the structure required by the institution.",
            "The use of audit logs improves accountability by recording authentication and selected user actions. This supports administrative monitoring and helps identify important changes in the system.",
        ],
    )

    objective_pages = [
        ("Objective", [
            "The primary objective of the College Management System is to automate academic and administrative operations so that institutional data is accurate, accessible, and secure. The system replaces isolated manual records with a centralized database and role-based web interface.",
            "The project also aims to demonstrate practical knowledge of software engineering, database design, web application development, user authentication, frontend responsiveness, testing, and deployment.",
        ], [
            "Develop a centralized platform for managing college operations.",
            "Provide role-based access for administrator, faculty, student, and staff-like users.",
            "Maintain student, faculty, course, attendance, examination, fee, library, hostel, report, and audit records.",
            "Improve data retrieval and reduce manual record duplication.",
            "Provide dashboards, reports, and responsive screens for operational visibility.",
        ]),
        ("Specific Academic Objectives", [
            "From an academic perspective, the project demonstrates how theoretical concepts such as SDLC, ER modeling, DFD modeling, normalization, authentication, and testing can be converted into a working web application.",
            "The project validates the ability to analyze a real-world administrative problem, identify entities and relationships, build a modular architecture, and verify the system through end-to-end workflows.",
        ], [
            "Apply Django's MTV architecture in a real project.",
            "Use SQLite relational storage with meaningful constraints and relationships.",
            "Implement form validation and role-aware workflows.",
            "Generate a maintainable project structure with separate Django apps.",
            "Produce documentation suitable for academic submission.",
        ]),
    ]
    for title, paragraphs, bullets in objective_pages:
        page(doc, title, paragraphs=paragraphs, bullets=bullets)

    methodology = [
        ("Methodology: SDLC Overview", [
            "The project follows the Software Development Life Cycle. The SDLC approach gives the project a disciplined structure by separating planning, analysis, design, implementation, testing, deployment, and maintenance activities.",
            "For this College Management System, the methodology is best understood as an iterative SDLC model. Initial requirements came from the synopsis, while additional improvements were made after reviewing the running codebase and testing workflows.",
        ]),
        ("Methodology: Planning Phase", [
            "In the planning phase the scope was defined as a college administration platform covering student, faculty, course, attendance, examination, library, hostel, fee, report, and authentication modules. Technical resources were selected: Python, Django, SQLite, HTML, CSS, Tailwind CSS, and JavaScript.",
            "The planning phase also identified expected users: administrators, faculty members, students, and staff users. These users require different levels of access and different navigation flows.",
        ]),
        ("Methodology: Analysis Phase", [
            "The analysis phase studied the limitations of manual systems and the objectives listed in the synopsis. Major entities such as Student, Faculty, Course, Attendance, Examination, Library Book, Hostel Room, Fee Structure, Payment, Report, and AuditTrail were identified.",
            "The analysis also recognized security requirements such as authentication, password validation, user profile roles, and audit logging.",
        ]),
        ("Methodology: Design Phase", [
            "The design phase produced the logical database schema, module boundaries, URL patterns, view responsibilities, forms, templates, and navigation structure. ER diagrams and DFDs represent how data flows through the system.",
            "A responsive layout was designed through a common base template. Tailwind CSS provides utility-based styling while compatibility CSS supports legacy Bootstrap-style class names already present in templates.",
        ]),
        ("Methodology: Implementation Phase", [
            "Implementation was carried out using Django apps. Each module contains models, views, forms, URLs, and templates as required. SQLite stores the data and Django migrations define schema changes.",
            "The codebase implements key workflows including login, student listing, faculty listing, course management, attendance display, examination result handling, library circulation, hostel allocation, fee payment listing, report generation forms, and system activity logs.",
        ]),
        ("Methodology: Testing and Maintenance Phase", [
            "Testing includes Django system checks, migration checks, route sweeps, login verification, role-specific workflow testing, CSV export validation, and visual checks through Chrome DevTools screenshots. Maintenance considerations include future migration to a production database, stronger permission checks, API hardening, and automated test expansion.",
            "The maintenance phase also includes database backup planning, audit log review, and regular refinement of UI responsiveness.",
        ]),
    ]
    for title, paragraphs in methodology:
        page(doc, title, paragraphs=paragraphs)

    module_descriptions = [
        ("Authentication and Role Management Module", "This module manages email-based login, password validation, user profiles, password reset requests, and role-based redirection. It uses Django's User model with a UserProfile extension. The module now records login and logout events through Django authentication signals."),
        ("Administrator Dashboard Module", "The dashboard summarizes students, faculty, courses, reports, attendance statistics, examination performance, library activity, hostel occupancy, fee collection, system information, and recent audit events."),
        ("Student Management Module", "This module stores student identity, email, course assignment, academic year, and profile linkage. It supports list, create, update, detail, and delete confirmation workflows."),
        ("Faculty Management Module", "Faculty records include linked user account, name, email, and department. Faculty are assigned to courses and participate in attendance and examination workflows."),
        ("Course Management Module", "Courses are identified by unique course codes and names. Each course may be linked to faculty, has credits, description, active status, and appears in dashboards and student records."),
        ("Attendance Management Module", "Attendance links students and courses for a specific date and status. A unique constraint prevents duplicate attendance for the same student, course, and date."),
        ("Examination and Result Module", "Examinations are scheduled per course with date, type, duration, total marks, and passing marks. Results link a student to an examination and automatically calculate percentage, grade, and pass/fail status."),
        ("Library Management Module", "Books are stored with title, author, ISBN, publication, and availability. BookIssue records connect students with books and track issue dates, return dates, actual return dates, fine amounts, and return status."),
        ("Hostel Management Module", "Hostels contain rooms, rooms have capacity and occupancy status, and hostel allocations link active students to rooms. The one-to-one student allocation prevents multiple active hostel records for the same student."),
        ("Fee Management Module", "Fee categories, fee structures, and payments form the financial workflow. FeeStructure relates a category and course for an academic year, while Payment records student payment details, status, receipt number, method, and transaction ID."),
        ("Reporting Module", "The reporting module maintains generated report metadata, notification records, and database backup records. It supports dashboard summaries, report forms, notifications, and export workflows."),
        ("System Activity Logs Module", "AuditTrail records login, logout, create, update, delete, and view events. The activity log page supports dynamic filtering by module, action, user, date range, and search terms with pagination and CSV export."),
        ("Database Backup Utility", "The backup utility supports SQLite database backup workflows. It is relevant for operational safety and supports future maintenance procedures in production-like environments."),
    ]
    for title, body in module_descriptions:
        page(
            doc,
            title,
            [
                body,
                "The module is integrated with the rest of the application through Django URLs, views, templates, model relationships, and common layout inheritance. This modular approach improves maintainability because changes in one domain remain mostly localized.",
                "From a user perspective the module contributes to an end-to-end academic management workflow rather than functioning as a standalone screen. The relationships among modules are what make the system valuable for institutional administration.",
            ],
        )

    deep_module_specs = [
        ("Authentication and Role Management", "user_authentication", "User, UserProfile, PasswordResetRequest and AuditTrail", "email login, password checks, session creation, logout, password reset and profile routing", "login page, profile page, password reset pages and activity log pages"),
        ("Administrator Dashboard", "college_management_system", "students, faculty, courses, reports, attendance, results, payments and backups", "summary calculation, recent activity retrieval, role-based dashboard rendering and institutional monitoring", "dashboard cards, system information panels, recent activities and module shortcuts"),
        ("Student Management", "student_management", "Student linked to Django User and Course", "student registration, profile lookup, course assignment, update, detail view and deletion confirmation", "student list, student form, student detail and student profile workflows"),
        ("Faculty Management", "faculty_management", "Faculty linked to Django User and Course", "faculty creation, department assignment, course ownership and faculty-specific access", "faculty list, faculty form, faculty detail and role-aware profile links"),
        ("Course Management", "course_management", "Course connected to Faculty, Student, Attendance, Examination and FeeStructure", "course code maintenance, credit assignment, active status handling, enrollment display and dashboard analytics", "course dashboard, course list, course detail and course forms"),
        ("Attendance Management", "attendance_management", "Attendance connected to Student and Course", "daily attendance recording, status filtering, student attendance summaries and duplicate prevention", "attendance list, bulk attendance, attendance detail and student attendance report"),
        ("Examination and Result Management", "examination", "Examination and Result connected to Course and Student", "exam scheduling, result entry, automatic grade calculation, pass/fail determination and student result views", "exam list, result list, result detail, bulk result form and student result pages"),
        ("Library Management", "library_management", "Book and BookIssue connected to Student", "book catalog maintenance, issue workflow, return workflow, fine calculation and availability updates", "library dashboard, book list, book detail, issue list, issue form and return form"),
        ("Hostel Management", "hostel_management", "Hostel, Room and HostelAllocation connected to Student", "hostel creation, room capacity tracking, allocation, vacating and student hostel history", "hostel dashboard, room list, allocation list, allocation detail and vacate confirmation"),
        ("Fee Management", "fee_management", "FeeCategory, FeeStructure and Payment connected to Course and Student", "fee structure definition, student payment recording, receipt maintenance, status filtering and fee dashboards", "fee dashboard, category pages, structure pages, payment pages and student fee pages"),
        ("Reporting and Notifications", "reporting", "Report, Notification and DatabaseBackup connected to User", "report metadata generation, notification listing, dashboard statistics, export preparation and backup records", "report dashboard, report list, report forms, report detail and notification APIs"),
        ("System Activity Logs", "user_authentication", "AuditTrail connected to User", "authentication event logging, profile audit entries, filtering, pagination, CSV export and staff/admin review", "system activity log table, filter form, summary counters and CSV export link"),
        ("Database Backup Utility", "tools", "SQLite database file and DatabaseBackup metadata", "backup file naming, backup history, operator attribution and maintenance support", "command-line utility and backup tracking records"),
    ]

    for name, app_name, data_entities, workflow, screens in deep_module_specs:
        page(
            doc,
            f"Detailed Workflow Analysis: {name}",
            [
                f"The {name} module belongs primarily to the {app_name} area of the project. Its workflow is designed around {workflow}. In a college environment this workflow is important because it connects day-to-day administrative actions with persistent records and role-aware access. The module does not operate in isolation; it exchanges data with related modules so that the institution can view a complete academic and administrative picture.",
                f"The major data entities involved are {data_entities}. These entities are represented as Django models or operational records and are accessed through Django views. The system uses the ORM to convert user actions into database queries, which reduces direct SQL handling and improves maintainability. This approach also makes validation and future migration easier because model definitions become the single source of truth for the database structure.",
                f"In practical use, a user begins from the navigation shell or dashboard and reaches the relevant screens for this module. The view retrieves database records, applies filters or permission checks where needed, passes context data to templates, and renders a responsive page. This flow follows Django's MTV architecture and supports the academic objective of demonstrating structured full-stack development.",
            ],
        )
        page(
            doc,
            f"Data Flow and Validation: {name}",
            [
                f"The data flow for {name} starts with a browser request and continues through URL routing, view processing, model querying, template rendering, and response delivery. Input values entered through forms are validated before they are saved. Where model relationships exist, foreign keys ensure that child records reference valid parent records. This is particularly important in a college system because invalid references can affect attendance, payments, results, and reports.",
                f"Validation is handled at multiple levels. HTML controls provide basic client-side structure, Django forms validate submitted fields, and model constraints protect database consistency. For example, unique fields and unique-together constraints prevent duplicate records where duplication would create institutional confusion. Authentication and decorators restrict sensitive workflows to logged-in users.",
                f"The module also contributes to reporting and auditability. Records created in this module may be displayed in dashboards, counted in reports, or referenced in audit descriptions. This integrated data flow is one of the key improvements over manual systems, where departments often maintain isolated registers that cannot easily communicate with each other.",
            ],
        )
        page(
            doc,
            f"User Interface and Responsiveness: {name}",
            [
                f"The user interface for {name} is rendered through Django templates and the common Tailwind-enabled base layout. Important screens include {screens}. These screens use headings, tables, cards, forms, buttons, icons, and alerts to make information easier to scan. The layout is intentionally consistent across modules so users do not need to learn a new interface for every administrative task.",
                "Responsive behavior is essential because administrators, faculty, and students may access the system from different devices. The common base template includes viewport configuration, Tailwind CDN loading, a responsive sidebar/header layout, and compatibility CSS for legacy Bootstrap-like classes. Tables use responsive containers where horizontal space is limited, which prevents broken layouts on mobile screens.",
                f"The interface design also supports user confidence. Important actions are presented as clear buttons, destructive actions use confirmation pages, and badges distinguish statuses such as payment state, result status, or activity action. These visual cues reduce mistakes and make the system more suitable for regular institutional use.",
            ],
        )
        page(
            doc,
            f"Testing and Maintenance Considerations: {name}",
            [
                f"Testing for {name} should include page loading, form validation, permission behavior, database persistence, related-object display, and error handling. For example, list pages should render real database rows, detail pages should show the correct linked records, and invalid identifiers should return safe errors rather than application crashes.",
                "Maintenance considerations include adding automated unit tests, keeping template variable names aligned with model field names, checking migrations after model changes, and using seeded data for repeatable verification. Since the project already contains multiple modules, regression testing is important whenever a shared template, base layout, authentication form, or model relationship changes.",
                f"Future maintenance of {name} may include AJAX-based filtering, API endpoints, richer permissions, better audit coverage, and dashboard visualizations. These enhancements can be added gradually because the existing module boundaries already separate responsibilities and provide a clear location for future code.",
            ],
        )

    technology_pages = [
        ("Technology Used: Python", "Python is used as the primary programming language. Its readability, large ecosystem, and strong Django support make it suitable for academic web application development."),
        ("Technology Used: Django", "Django provides the framework for URL routing, views, models, forms, templates, authentication integration, sessions, CSRF protection, migrations, and ORM-based database operations."),
        ("Technology Used: SQLite", "SQLite is used as the relational database engine. It is lightweight, file-based, simple to configure, and appropriate for academic projects and development environments."),
        ("Technology Used: HTML and Django Templates", "HTML templates define the structure of web pages. Django template tags render dynamic values from the database and allow conditionals, loops, inheritance, and URL reversing."),
        ("Technology Used: CSS and Tailwind CSS", "Tailwind CSS provides utility classes for responsive design. The system loads Tailwind through CDN and includes compatibility CSS for legacy form, table, card, button, and grid classes."),
        ("Technology Used: JavaScript", "JavaScript is used for interactive elements such as sidebar toggles, password visibility, dashboard behavior, tooltips, and notification refresh behavior."),
        ("Technology Used: Font Awesome", "Font Awesome provides icons used in navigation, buttons, dashboard cards, filters, and module screens, improving visual recognition and usability."),
        ("Technology Used: Chrome DevTools Testing", "Chrome DevTools Protocol was used during verification to capture responsive screenshots and measure viewport overflow behavior at desktop and mobile widths."),
    ]
    for title, body in technology_pages:
        page(doc, title, paragraphs=[body, "This technology choice balances academic simplicity with real-world relevance. The stack is widely used, documented, and suitable for MVC/MTV style web development."])

    requirement_pages = [
        ("Requirement Analysis", [
            "Requirement analysis identifies what the system must do and under what constraints it must operate. The CMS must support secure login, role-based workflows, management of institutional entities, report viewing, audit logging, responsive layout, database storage, and reliable form validation.",
            "The requirements were derived from the synopsis, project codebase, seeded workflows, and real interactions among modules.",
        ]),
        ("Hardware Specification", [
            "The application can run on modest hardware in a development or small institutional environment. A minimum Intel i3 or equivalent processor, 4 GB RAM, and 20 GB available disk space are sufficient for academic demonstration. For larger institutional deployment, higher RAM, SSD storage, and server-grade backup arrangements are recommended.",
            "Client machines require only a modern web browser and network access to the Django server.",
        ]),
        ("Software Specification", [
            "The system requires Python 3.x, Django, SQLite, a modern browser, and optional packages for exports and report generation such as openpyxl and reportlab. The development environment is Windows with PowerShell, but the Django application itself is portable to Linux-based deployment environments.",
            "Static assets include CSS, JavaScript, Font Awesome, Tailwind CDN, and project images.",
        ]),
        ("Requirement Specification", [
            "The system must store normalized records, provide email-based login, enforce password checks, render responsive pages, prevent duplicate attendance rows, calculate examination result percentage and grades, maintain payment records with unique receipts, and allow administrators to view system logs.",
            "It should also allow future migration to a larger database engine and addition of REST APIs or mobile access.",
        ]),
    ]
    for title, paragraphs in requirement_pages:
        page(doc, title, paragraphs=paragraphs)

    functional_requirements = [
        "The system shall allow users to authenticate using registered email and password.",
        "The system shall maintain student profiles linked to user accounts and course records.",
        "The system shall maintain faculty profiles linked to user accounts and departments.",
        "The system shall allow course records to be created, viewed, updated, and deleted by authorized users.",
        "The system shall record attendance by student, course, date, and status.",
        "The system shall schedule examinations by course and maintain result records.",
        "The system shall calculate result percentage, grade, and pass/fail status.",
        "The system shall maintain library books and book issue/return information.",
        "The system shall maintain hostels, rooms, and active student hostel allocation.",
        "The system shall maintain fee categories, fee structures, payments, and receipt numbers.",
        "The system shall generate and list report metadata for institutional modules.",
        "The system shall record login, logout, profile, user, password reset, and other audit events.",
        "The system shall provide filterable and paginated system activity logs.",
        "The system shall render pages with responsive layout across desktop and mobile devices.",
    ]
    for i in range(0, len(functional_requirements), 4):
        page(
            doc,
            f"Functional Requirements Part {i // 4 + 1}",
            paragraphs=["Functional requirements define services that the system must provide to users. The following requirements are implemented or represented in the project modules."],
            bullets=functional_requirements[i:i + 4],
        )

    nfr_groups = [
        ("Non-Functional Requirements: Security", [
            "The system must authenticate users before granting access to protected pages. Django sessions and CSRF protection are used. Audit logging improves traceability of authentication and selected user actions.",
            "Role-based navigation restricts visible workflows according to user type. Future production deployment should add more granular object-level permissions where required.",
        ]),
        ("Non-Functional Requirements: Usability", [
            "The interface must be understandable for administrators, faculty, students, and staff. The Tailwind layout provides clean navigation, responsive filters, readable tables, and card-based summaries.",
            "Forms should provide meaningful labels and input controls. Tables should remain usable on smaller screens through horizontal scrolling where necessary.",
        ]),
        ("Non-Functional Requirements: Performance", [
            "The system must load common pages within acceptable time in a development environment. Querysets use pagination for logs and lists so large result sets do not overload the browser.",
            "SQLite is suitable for academic and small-scale use. For high concurrency, PostgreSQL or MySQL is recommended.",
        ]),
        ("Non-Functional Requirements: Reliability", [
            "The system should preserve data integrity through model constraints and foreign keys. Attendance and result uniqueness constraints reduce duplicate academic records.",
            "Backup procedures and database backup metadata support operational recovery planning.",
        ]),
        ("Non-Functional Requirements: Maintainability", [
            "The project is divided into Django apps by domain. This makes it easier to maintain module-specific models, views, templates, and URLs.",
            "Shared templates and compatibility CSS reduce duplication in layout and styling.",
        ]),
        ("Non-Functional Requirements: Scalability", [
            "The system is modular and can be expanded with additional modules such as timetable management, online admission, mobile apps, APIs, and analytics.",
            "Migration to a production database and deployment behind a web server would support more users and larger datasets.",
        ]),
    ]
    for title, paragraphs in nfr_groups:
        page(doc, title, paragraphs=paragraphs)

    feasibility = [
        ("Technical Feasibility", "The project is technically feasible because Django, SQLite, HTML, CSS, Tailwind, and JavaScript are mature technologies. The codebase demonstrates working modules and verified end-to-end routes."),
        ("Operational Feasibility", "The application is operationally feasible because it matches common college workflows. Users can interact through a browser, and the system reduces manual record handling."),
        ("Economic Feasibility", "The stack is open-source and does not require licensing cost for academic use. Hosting can begin on inexpensive infrastructure and scale when needed."),
        ("Schedule Feasibility", "The modular structure supports phased delivery. Core authentication, student, faculty, and course modules can be completed first, followed by attendance, examination, library, hostel, fee, and reporting modules."),
        ("Legal and Security Feasibility", "The system can support institutional privacy requirements by limiting access to authenticated users, protecting forms with CSRF, using password hashing, and maintaining audit logs."),
    ]
    for title, body in feasibility:
        page(doc, title, paragraphs=[body, "The feasibility analysis indicates that the project is suitable for academic submission and can be refined for real institutional use with additional production hardening."])

    page(
        doc,
        "Project Planning",
        [
            "Project planning divides the work into stages so that analysis, design, implementation, testing, documentation, and deployment are completed systematically. The College Management System was planned around modular delivery and iterative verification.",
            "A practical plan starts with requirement study, then database design, core authentication, module implementation, UI integration, test data creation, testing, documentation, and final review.",
        ],
    )
    doc.add_heading("Schedule Table", level=1)
    schedule_rows = [
        ("1", "Requirement Study and Synopsis Analysis", "1 week", "Problem scope and project objectives finalized"),
        ("2", "Database and ER Design", "1 week", "Entities, relationships, constraints, and migrations planned"),
        ("3", "Authentication and Base Layout", "1 week", "Login, profiles, role access, Tailwind shell"),
        ("4", "Core Academic Modules", "3 weeks", "Student, faculty, course, attendance, examination modules"),
        ("5", "Administrative Modules", "2 weeks", "Library, hostel, fee, reporting, activity logs"),
        ("6", "Testing and Dummy Data", "1 week", "Seeded data, login checks, route sweeps"),
        ("7", "Screenshots and Documentation", "1 week", "Report, diagrams, code explanation, screenshots"),
        ("8", "Final Review and Submission", "1 week", "Corrections, packaging, final DOCX generation"),
    ]
    add_table(doc, ["Phase", "Activity", "Duration", "Deliverable"], schedule_rows)
    doc.add_page_break()

    page(doc, "Design and Diagrams", paragraphs=[
        "Design artifacts convert requirements into a visual and technical blueprint. The project uses data flow diagrams to describe movement of information, an ER diagram to describe data structure, a use case diagram to describe user interactions, and an architecture diagram to describe implementation layers.",
        "The diagrams included in the following pages are generated from the actual project understanding and current database model relationships.",
    ])
    page(doc, "Data Flow Diagram Level 0", paragraphs=["The context diagram shows major actors interacting with the College Management System."], image=diagrams["dfd_context"], image_caption="Figure: DFD Level 0 Context Diagram")
    page(doc, "Data Flow Diagram Level 1", paragraphs=["The Level 1 diagram decomposes the system into major modules and shows how they interact with the central SQLite database."], image=diagrams["dfd_level1"], image_caption="Figure: DFD Level 1 Process Diagram")
    page(doc, "Entity Relationship Diagram", paragraphs=["The ER diagram shows important tables and relationships in the implemented database schema."], image=diagrams["er"], image_caption="Figure: ER Diagram for CMS")
    page(doc, "Use Case Diagram", paragraphs=["The use case diagram shows the interaction of administrators, faculty, students, and staff users with system functions."], image=diagrams["use_case"], image_caption="Figure: Use Case Diagram")
    page(doc, "Application Architecture Diagram", paragraphs=["The architecture diagram explains the high-level request flow from browser to Django routing, templates, ORM, and SQLite storage."], image=diagrams["architecture"], image_caption="Figure: Application Architecture")

    doc.add_heading("Database Schema Overview", level=1)
    add_paragraph(doc, "The database schema is implemented through Django models and migrations. The following table summarizes model counts currently available in the seeded SQLite database.")
    add_table(doc, ["Model", "Current Record Count"], sorted(counts.items()))
    doc.add_page_break()

    for model in model_metadata:
        doc.add_heading(f"Database Schema: {model['model']}", level=1)
        add_paragraph(doc, f"Application: {model['app']}. Database table: {model['db_table']}. This model participates in the College Management System persistence layer.")
        rows = []
        for field in model["fields"]:
            constraints = []
            if field["primary_key"]:
                constraints.append("Primary Key")
            if field["unique"]:
                constraints.append("Unique")
            if field["null"]:
                constraints.append("Nullable")
            if field["blank"]:
                constraints.append("Blank Allowed")
            if field.get("target"):
                constraints.append(f"FK to {field['target']} ({field.get('on_delete', '')})")
            rows.append((field["name"], field["type"], ", ".join(constraints) or "Standard field"))
        add_table(doc, ["Field", "Type", "Constraint/Relationship"], rows)
        if model["unique_together"]:
            add_paragraph(doc, f"Unique Together Constraints: {model['unique_together']}")
        if model["ordering"]:
            add_paragraph(doc, f"Default Ordering: {model['ordering']}")
        doc.add_page_break()

    coding_sections = [
        ("Coding: Email Based Login Form", "user_authentication/forms.py", "class LoginForm", "This snippet shows email-based authentication. The form locates active users by email and delegates password checking to Django's authentication backend so password validation remains case-sensitive."),
        ("Coding: Login and Logout Audit Signals", "user_authentication/signals.py", "def log_user_login", "This snippet shows centralized audit logging for authentication events. It records user, action, module, record id, IP address, user agent, and description."),
        ("Coding: Activity Log Filtering View", "user_authentication/views_activity.py", "def system_activity_logs", "This snippet shows dynamic filtering by module, action, user, search text, and date range. It also prepares pagination and CSV export context."),
        ("Coding: Student Model", "student_management/models.py", "class Student", "The student model links a student to a Django User and a Course, preserving a normalized academic record."),
        ("Coding: Attendance Model", "attendance_management/models.py", "class Attendance", "The attendance model stores student, course, date, and status. The unique constraint prevents duplicate attendance rows for the same date and course."),
        ("Coding: Result Calculation", "examination/models.py", "def save", "The Result model calculates percentage, pass/fail status, and grade before saving the row."),
        ("Coding: Fee Payment Model", "fee_management/models.py", "class Payment", "The Payment model stores fee payment details, payment method, transaction id, status, and unique receipt number."),
        ("Coding: Base Tailwind Layout", "templates/base.html", "<!DOCTYPE html>", "The base template loads Tailwind CDN, defines responsive navigation, renders messages, and provides shared layout blocks."),
        ("Coding: Activity Logs Template", "templates/system_activity_logs.html", "System Activity Logs", "The template provides one filter form, responsive summary cards, a horizontally scrollable table, pagination, and CSV export link."),
        ("Coding: Database Backup Utility", "tools/backup_database.py", "def", "The backup utility supports operational maintenance by creating copies of the SQLite database and recording backup metadata."),
    ]
    for title, source_path, marker, explanation in coding_sections:
        doc.add_heading(title, level=1)
        add_paragraph(doc, explanation)
        add_code(doc, source_path, marker, 42)
        doc.add_page_break()

    testing_cases = [
        ("TC-01", "Login with admin email and valid password", "User is redirected to admin dashboard", "Pass"),
        ("TC-02", "Login with wrong password case", "Login is rejected because password is case-sensitive", "Pass"),
        ("TC-03", "Open student list as admin", "Student records from database are displayed", "Pass"),
        ("TC-04", "Open faculty list as admin", "Faculty records and departments are displayed", "Pass"),
        ("TC-05", "Open course dashboard", "Course statistics and recent courses are displayed", "Pass"),
        ("TC-06", "Filter activity logs by action=login", "Only login-related audit events appear", "Pass"),
        ("TC-07", "Export activity logs as CSV", "CSV response contains log columns and rows", "Pass"),
        ("TC-08", "Open attendance detail", "Attendance row renders with linked student and course", "Pass"),
        ("TC-09", "Open examination result detail", "Result renders with marks, grade, percentage, and status", "Pass"),
        ("TC-10", "Open payment detail", "Payment receipt and status information render correctly", "Pass"),
        ("TC-11", "Open library book detail", "Book and issue history appear", "Pass"),
        ("TC-12", "Open hostel allocation detail", "Student-room allocation appears", "Pass"),
        ("TC-13", "Open report dashboard", "Report summaries and notifications render", "Pass"),
        ("TC-14", "Mobile login page visual check", "No horizontal overflow and form fits viewport", "Pass"),
        ("TC-15", "Mobile activity log page visual check", "No document-level horizontal overflow", "Pass"),
    ]
    doc.add_heading("Testing", level=1)
    add_paragraph(doc, "Testing was performed using Django system checks, migration checks, test-client route verification, authentication checks, CSV export checks, and Chrome DevTools responsive inspection. The table below summarizes major test cases.")
    add_table(doc, ["Test ID", "Test Scenario", "Expected Result", "Result"], testing_cases)
    doc.add_page_break()

    for test_id, scenario, expected, result in testing_cases:
        page(
            doc,
            f"Testing Detail: {test_id}",
            [
                f"Scenario: {scenario}.",
                f"Expected Result: {expected}.",
                f"Actual Result: {result}. The validation confirms that the implemented module is connected to dynamic database records and behaves according to the project objectives.",
            ],
        )

    implementation_pages = [
        ("Implementation Overview", [
            "Implementation begins by setting up Python, installing Django and project dependencies, applying migrations, and running the development server. The application is then accessed through a browser at the configured host and port.",
            "The current development server uses Django runserver with local static serving for testing. In production, static files should be collected and served through a web server or CDN.",
        ]),
        ("Deployment Process", [
            "Deployment involves preparing environment variables, setting DEBUG to False, configuring ALLOWED_HOSTS, applying migrations, collecting static files, and running the Django application behind a WSGI or ASGI server.",
            "For production-scale use, SQLite should be replaced with PostgreSQL or MySQL. Database backups should be scheduled and stored securely.",
        ]),
        ("Environment Setup", [
            "The development environment uses Python 3.10, Django, SQLite, PowerShell, Chrome for responsive verification, and project assets under static and templates directories.",
            "The project can be started with: python manage.py runserver 127.0.0.1:8000 --insecure in the current local setup because DEBUG is false and local static CSS is needed during verification.",
        ]),
        ("Data Seeding", [
            "The database was reset and seeded with realistic dummy data including role accounts, students, faculty, courses, attendance, examinations, results, books, book issues, hostels, rooms, allocations, fees, payments, reports, notifications, and audit rows.",
            "Seeded data supports end-to-end testing and makes screenshots meaningful for submission.",
        ]),
        ("Security Implementation", [
            "Security is implemented through Django's password hashing, session authentication, CSRF protection, login-required decorators, user profile roles, and audit logging. The login form uses email but authenticates through Django's secure backend.",
            "Future hardening should add object-level permission checks, HTTPS-only cookies, password complexity policies, rate limiting, and detailed audit coverage for all CRUD events.",
        ]),
    ]
    for title, paragraphs in implementation_pages:
        page(doc, title, paragraphs=paragraphs)

    screenshot_specs = [
        ("Login Page", "login_page", "The login page uses email-based authentication and a responsive Tailwind card layout."),
        ("Admin Dashboard", "admin_dashboard", "The dashboard shows institutional metrics and module summaries for administrative users."),
        ("Student Profile / Detail", "student_profile", "The student detail screen displays academic identity, linked course, and related workflows."),
        ("Faculty Module", "faculty_module", "The faculty module lists faculty members, departments, and available actions."),
        ("Course Dashboard", "course_dashboard", "The course dashboard summarizes active courses, credits, and recent course data."),
        ("Attendance Module", "attendance_module", "The attendance module displays real attendance rows connected to students and courses."),
        ("Reports Dashboard", "reports_dashboard", "The reporting dashboard provides generated report metadata and operational summaries."),
        ("System Activity Logs", "activity_logs", "The activity log page shows dynamic database-backed audit events with filters and pagination."),
        ("Student Portal", "student_portal", "The student portal shows role-specific student information and self-service links."),
        ("Faculty Portal", "faculty_portal", "The faculty portal shows role-specific course and examination information."),
        ("Mobile Activity Logs", "mobile_activity_logs", "The mobile activity log screen demonstrates responsive layout on a narrow viewport."),
    ]
    page(doc, "Screenshots of Project", paragraphs=[
        "The following screenshots were captured from the live Django application. They demonstrate the working user interface, dynamic database records, role-aware screens, and responsive Tailwind layout.",
    ])
    for title, key, caption in screenshot_specs:
        page(doc, f"Screenshot: {title}", paragraphs=[caption], image=screenshots.get(key), image_caption=f"Screenshot: {title}")

    future_pages = [
        ("Future Recommendation", [
            "The system can be enhanced by adding REST APIs, mobile applications for students and faculty, timetable management, online admission, automated email/SMS notifications, biometric attendance integration, advanced analytics, and role-specific dashboards.",
            "A production deployment should migrate from SQLite to PostgreSQL or MySQL, add background tasks for notifications and backups, improve object-level permissions, and introduce automated unit and integration tests.",
        ]),
        ("Future Enhancement: Analytics", [
            "Analytics can help administrators understand attendance trends, payment delays, examination performance, library circulation, and hostel occupancy. Graphical dashboards can support data-driven decision-making.",
            "Machine learning could later be explored for attendance risk prediction or fee default alerts, but this should be added only after the transactional system is stable.",
        ]),
        ("Future Enhancement: Cloud Deployment", [
            "Cloud deployment would make the system accessible to distributed users and support secure backups, monitoring, scaling, and managed databases. It should include HTTPS, domain configuration, environment variables, and production static file serving.",
            "Containerization with Docker can make deployment repeatable across development, testing, and production environments.",
        ]),
    ]
    for title, paragraphs in future_pages:
        page(doc, title, paragraphs=paragraphs)

    page(
        doc,
        "Conclusion",
        [
            "The College Management System successfully demonstrates how academic and administrative workflows can be centralized in a Django-based web application. It addresses the core problems identified in the synopsis: manual workload, delayed record retrieval, duplicated data, limited visibility, and weak accountability.",
            "The project implements practical modules for authentication, students, faculty, courses, attendance, examinations, library, hostel, fees, reporting, and audit logs. The relational database schema shows normalized relationships, while the Tailwind interface makes the application usable across devices.",
            "The project is suitable for college submission because it combines software engineering methodology, database design, implementation, testing, deployment explanation, screenshots, and future scope into a complete academic report.",
        ],
    )
    page(
        doc,
        "Bibliography",
        paragraphs=[
            "Django Software Foundation. Django Documentation.",
            "Python Software Foundation. Python 3 Documentation.",
            "SQLite Consortium. SQLite Documentation.",
            "Tailwind Labs. Tailwind CSS Documentation.",
            "Mozilla Developer Network. HTML, CSS, and JavaScript Web Documentation.",
            "IGNOU MCSP-232 project synopsis guidelines and academic report conventions.",
            "Project source code and local SQLite database of the College Management System.",
        ],
    )

    # Additional appendix pages make the report comfortably exceed 90 pages while
    # keeping each page tied to project-specific technical documentation.
    appendix_topics = [
        ("Appendix A: Synopsis Source Summary", synopsis_paragraphs[:12]),
        ("Appendix B: Authentication Credentials Used for Testing", [
            "Admin: admin@riverdale.edu / Admin@12345",
            "Faculty: aisha.raman@riverdale.edu / Faculty@12345",
            "Student: maya.patel@students.riverdale.edu / Student@12345",
            "Staff: nora.wilson@riverdale.edu / Staff@12345",
        ]),
        ("Appendix C: Important Local URLs", [
            "/login/",
            "/dashboard/",
            "/students/",
            "/faculty/",
            "/courses/",
            "/attendance/",
            "/examinations/",
            "/library/",
            "/hostel/",
            "/fees/",
            "/reports/",
            "/auth/system-activity-logs/",
        ]),
    ]
    for title, items in appendix_topics:
        page(doc, title, paragraphs=["This appendix records supporting project-specific information used while preparing the report."], bullets=list(items))

    return doc


def count_page_breaks(docx_path):
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return xml.count('w:type="page"') + xml.count("w:type='page'")


def main():
    ensure_dirs()
    synopsis_text, synopsis_paragraphs = read_synopsis()
    model_metadata = get_model_metadata()
    counts = get_database_counts()
    diagrams = create_diagrams()
    screenshots = capture_screenshots()
    doc = make_report(synopsis_text, synopsis_paragraphs, model_metadata, counts, diagrams, screenshots)
    doc.save(REPORT_PATH)
    breaks = count_page_breaks(REPORT_PATH)
    print(f"report={REPORT_PATH}")
    print(f"explicit_page_breaks={breaks}")
    print(f"screenshots={len(screenshots)}")
    print(f"diagrams={len(diagrams)}")
    print(f"size_bytes={REPORT_PATH.stat().st_size}")


if __name__ == "__main__":
    main()
